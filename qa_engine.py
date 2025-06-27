from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.llms import Ollama
from langchain.chains.question_answering import load_qa_chain
from docx import Document
from langchain.schema import Document as LangchainDocument

# ✅ For PDFs
def process_pdf(file_path):
    loader = PyPDFLoader(file_path)
    pages = loader.load()
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    docs = text_splitter.split_documents(pages)

    embeddings = OllamaEmbeddings(model="llama3")
    vector_store = FAISS.from_documents(docs, embeddings)
    return vector_store

# ✅ For Word (.docx)
def process_docx(file_path):
    doc = Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    lc_doc = LangchainDocument(page_content=full_text, metadata={"source": file_path})

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    docs = text_splitter.split_documents([lc_doc])

    embeddings = OllamaEmbeddings(model="llama3")
    vector_store = FAISS.from_documents(docs, embeddings)
    return vector_store

# ✅ Query Answering
def answer_query(vector_store, query):
    llm = Ollama(model="llama3")
    retriever = vector_store.as_retriever()
    docs = retriever.get_relevant_documents(query)

    chain = load_qa_chain(llm, chain_type="stuff")
    response = chain.run(input_documents=docs, question=query)
    return response
